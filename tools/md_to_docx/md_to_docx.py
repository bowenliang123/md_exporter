import io
import logging
import re
from typing import Generator

import markdown
from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.shared import RGBColor
from htmldocx import HtmlToDocx

from tools.md_to_docx.font_enum import DocxFontEnum
from tools.utils.file_utils import get_meta_data
from tools.utils.mimetype_utils import MimeType
from tools.utils.param_utils import get_md_text


class MarkdownToDocxTool(Tool):
    logger = logging.getLogger(__name__)

    def _invoke(self, tool_parameters: dict) -> Generator[ToolInvokeMessage, None, None]:
        """
        invoke tools
        """
        # get parameters
        md_text = get_md_text(tool_parameters, is_strip_wrapper=True)
        try:
            html = markdown.markdown(text=md_text, extensions=["extra", "toc"])

            # transform html to docx
            new_parser = HtmlToDocx()
            doc: Document = new_parser.parse_html_string(html)

            # Set fonts for all text elements
            try:
                self.set_fonts_for_all_runs(doc)
                # Set table borders and text color
                self.set_table_border_and_text_color(doc)
            except Exception as e:
                self.logger.exception(e)

            result_bytes_io = io.BytesIO()
            doc.save(result_bytes_io)
            result_file_bytes = result_bytes_io.getvalue()
        except Exception as e:
            self.logger.exception("Failed to convert file")
            yield self.create_text_message(f"Failed to convert markdown text to DOCX file, error: {str(e)}")
            return

        yield self.create_blob_message(
            blob=result_file_bytes,
            meta=get_meta_data(
                mime_type=MimeType.DOCX,
                output_filename=tool_parameters.get("output_filename"),
            ),
        )
        return

    def is_contains_chinese_chars(self, text: str) -> bool:
        return bool(re.search(r"[一-鿿]", text))

    def set_chinese_fonts(self, doc):
        # Setting fonts globally
        # https://github.com/python-openxml/python-docx/issues/346#issuecomment-1698885586
        # https://zhuanlan.zhihu.com/p/548039429
        style = doc.styles["Normal"]
        font = style.font
        font.name = DocxFontEnum.TIMES_NEW_ROMAN
        rPr = style.element.get_or_add_rPr()
        rPr.rFonts.set(qn("w:eastAsia"), DocxFontEnum.SONG_TI)

    def set_fonts_for_all_runs(self, doc: Document):
        """Set Times New Roman for English text and SimSun for Chinese text in all text elements."""

        # Process all paragraphs in the document
        paragraph: Paragraph
        for paragraph in doc.paragraphs:
            run: Run
            for run in paragraph.runs:
                self.apply_fonts_to_run(run)

        # Process all paragraphs in tables
        table: Table
        for table in doc.tables:
            for row in table.rows:
                cell: _Cell
                for cell in row.cells:
                    paragraph: Paragraph
                    for paragraph in cell.paragraphs:
                        run: Run
                        for run in paragraph.runs:
                            self.apply_fonts_to_run(run)

    def apply_fonts_to_run(self, run: Run):
        if not run or not run.text:  # skip elements without text
            return

        try:
            # Set default font to Times New Roman
            run.font.name = DocxFontEnum.TIMES_NEW_ROMAN
            # Set East Asian font to SimSun
            run._element.rPr.rFonts.set(qn("w:eastAsia"), DocxFontEnum.SONG_TI)
            # Set ASCII font to Times New Roman
            run._element.rPr.rFonts.set(qn("w:ascii"), DocxFontEnum.TIMES_NEW_ROMAN)
            # Set high ANSI font to Times New Roman
            run._element.rPr.rFonts.set(qn("w:hAnsi"), DocxFontEnum.TIMES_NEW_ROMAN)
            # Set text color to black
            run.font.color.rgb = RGBColor(0, 0, 0)
        except Exception as e:
            self.logger.exception("Failed to apply fonts to run")

    def set_table_border_and_text_color(self, doc: Document):
        """Set table borders to solid lines and ensure all text is black."""
        
        # Process all tables
        for table in doc.tables:
            # Set table style to grid (solid lines)
            table.style = "Table Grid"
            
            # Process all cells
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text:  # Only process runs with text
                                run.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black
