import io
import logging
import re
from typing import Generator

import markdown
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.shared import RGBColor
from htmldocx import HtmlToDocx

# 配置日志
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

class DocxFontEnum:
    TIMES_NEW_ROMAN = "Times New Roman"
    SONG_TI = "SimSun"

class DocxConverter:
    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def convert(self, md_text: str, output_filename: str) -> bytes:
        """
        Convert markdown text to docx format
        """
        try:
            html = markdown.markdown(text=md_text, extensions=["extra", "toc"])

            # transform html to docx
            new_parser = HtmlToDocx()
            doc: Document = new_parser.parse_html_string(html)

            # Set fonts for all text elements
            try:
                self.set_fonts_for_all_runs(doc)
                # Set table borders and text color
                self.set_table_border_and_text_color(doc)
            except Exception as e:
                self.logger.exception(e)

            result_bytes_io = io.BytesIO()
            doc.save(result_bytes_io)
            return result_bytes_io.getvalue()
        except Exception as e:
            self.logger.exception("Failed to convert file")
            raise e

    def is_contains_chinese_chars(self, text: str) -> bool:
        return bool(re.search(r"[一-鿿]", text))

    def set_chinese_fonts(self, doc):
        # Setting fonts globally
        style = doc.styles["Normal"]
        font = style.font
        font.name = DocxFontEnum.TIMES_NEW_ROMAN
        rPr = style.element.get_or_add_rPr()
        rPr.rFonts.set(qn("w:eastAsia"), DocxFontEnum.SONG_TI)

    def set_fonts_for_all_runs(self, doc: Document):
        """Set Times New Roman for English text and SimSun for Chinese text in all text elements."""

        # Process all paragraphs in the document
        paragraph: Paragraph
        for paragraph in doc.paragraphs:
            run: Run
            for run in paragraph.runs:
                self.apply_fonts_to_run(run)

        # Process all paragraphs in tables
        table: Table
        for table in doc.tables:
            for row in table.rows:
                cell: _Cell
                for cell in row.cells:
                    paragraph: Paragraph
                    for paragraph in cell.paragraphs:
                        run: Run
                        for run in paragraph.runs:
                            self.apply_fonts_to_run(run)

    def apply_fonts_to_run(self, run: Run):
        if not run or not run.text:  # skip elements without text
            return

        try:
            # Set default font to Times New Roman
            run.font.name = DocxFontEnum.TIMES_NEW_ROMAN
            # Set East Asian font to SimSun
            run._element.rPr.rFonts.set(qn("w:eastAsia"), DocxFontEnum.SONG_TI)
            # Set ASCII font to Times New Roman
            run._element.rPr.rFonts.set(qn("w:ascii"), DocxFontEnum.TIMES_NEW_ROMAN)
            # Set high ANSI font to Times New Roman
            run._element.rPr.rFonts.set(qn("w:hAnsi"), DocxFontEnum.TIMES_NEW_ROMAN)
            # Set text color to black
            run.font.color.rgb = RGBColor(0, 0, 0)
        except Exception as e:
            self.logger.exception("Failed to apply fonts to run")

    def set_table_border_and_text_color(self, doc: Document):
        """Set table borders to solid lines and ensure all text is black."""
        
        # Process all tables
        for table in doc.tables:
            # Set table style to grid (solid lines)
            table.style = "Table Grid"
            
            # Process all cells
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text:  # Only process runs with text
                                run.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black

def test_conversion():
    try:
        # 读取测试文档
        logger.info("读取测试文档...")
        with open("test_document.md", "r", encoding="utf-8") as f:
            md_content = f.read()
        
        logger.info("创建转换工具实例...")
        # 创建转换工具实例
        converter = DocxConverter()
        
        logger.info("开始转换...")
        # 执行转换
        result_bytes = converter.convert(md_content, "test_output.docx")
        
        # 保存输出文件
        logger.info("保存输出文件...")
        with open("test_output.docx", "wb") as f:
            f.write(result_bytes)
        logger.info("转换完成，请查看 test_output.docx")
    except Exception as e:
        logger.exception("转换过程中出现错误:")

if __name__ == "__main__":
    test_conversion()
